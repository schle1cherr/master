# extractors.py – Extraktion und Segmentierung von Dokumentinhalten für KI-basiertes Retrieval
# Dieses Modul liefert Funktionen, um verschiedene Dokumenttypen (PDF, DOCX, XLSX) auszulesen, 
# sinnvoll zu segmentieren ("Chunking") und als normierte Dokumentobjekte bereitzustellen.
# Ziel: Verlustfreie, strukturierte Aufbereitung als Basis für das semantische und hybride Retrieval.

from pathlib import Path
import docx
import openpyxl
from PyPDF2 import PdfReader
from langchain_core.documents import Document
import re

from typing import List

def extract_text_from_pdf(file_path: Path) -> List[Document]:
    # Extrahiert Text aus PDF-Dateien.
    # Versucht zunächst direkte Textextraktion. 
    # (Optional: Fallback auf OCR kann bei gescannten PDFs ergänzt werden, wenn nötig.)
    try:
        reader = PdfReader(str(file_path))
        documents = []

        for page_num, page in enumerate(reader.pages):
            raw_text = page.extract_text() or ""  # Leere Zeichenkette, falls nichts gefunden wird
            lines = raw_text.splitlines()

            # Zusammenführung von Zeilen, um logische Abschnitte besser zu erkennen
            merged_lines = []
            buffer = ""
            for line in lines:
                clean = line.strip()
                if not clean:
                    continue
                # Annahme: Kleinbuchstabe am Zeilenanfang = Fortsetzung des vorherigen Satzes
                if re.match(r"^[a-zäöüß]", clean, re.IGNORECASE) and buffer:
                    buffer += " " + clean
                else:
                    if buffer:
                        merged_lines.append(buffer.strip())
                    buffer = clean
            if buffer:
                merged_lines.append(buffer.strip())

            # Segmentierung: Chunking entlang von Paragraphen (z. B. § 5, § 6 Abs. 2)
            page_text = "\n".join(line.strip() for line in lines if line.strip())

            paragraph_chunks = re.split(
                r"(?=\n?\s*§\s?\d+[a-zA-Z]?(?:\s*Abs\.\s*\d+)?\b)", 
                page_text
            )

            for chunk in paragraph_chunks:
                clean_chunk = chunk.strip()
                if not clean_chunk:
                    continue

                # Sicherstellen, dass das Chunk mit "§" beginnt (falls beim Split verloren)
                if not clean_chunk.startswith("§"):
                    clean_chunk = "§ " + clean_chunk

                match = re.search(r"§\s?(\d+[a-zA-Z]?)", clean_chunk)
                paragraph = match.group(1) if match else None

                # Dokumentenobjekt speichern, inkl. Metadaten für Rückverfolgung im Retrieval
                documents.append(
                    Document(
                        page_content=clean_chunk,
                        metadata={
                            "source": file_path.name,
                            "page_number": page_num + 1,
                            "paragraph": paragraph
                        }
                    )
                )

        return documents

    except Exception as e:
        # Fehlerausgabe für Debugging, keine Exception nach außen werfen
        print(f"Fehler beim Verarbeiten von PDF {file_path.name}: {e}")
        return []

def extract_text_from_docx(file_path: Path) -> List[Document]:
    # Extrahiert und sammelt den gesamten Text aus einem Word-Dokument (.docx)
    try:
        doc = docx.Document(str(file_path))
        # Nur nicht-leere Absätze verarbeiten
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [Document(page_content=full_text.strip(), metadata={"source": str(file_path)})]
    except Exception as e:
        print(f"Fehler beim DOCX {file_path.name}: {e}")
        return []

def extract_text_from_xlsx(file_path: Path) -> List[Document]:
    # Extrahiert Inhalte aus allen Zellen einer Excel-Datei, zeilenweise.
    try:
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        text = ""
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                # Nur belegte Zellen als Text speichern
                row_text = [str(cell) for cell in row if cell is not None]
                if row_text:
                    text += " | ".join(row_text) + "\n"
        return [Document(page_content=text.strip(), metadata={"source": str(file_path)})]
    except Exception as e:
        print(f"Fehler beim XLSX {file_path.name}: {e}")
        return []

def extract_all_documents_from_folder(folder_path: Path) -> List[Document]:
    # Lädt und extrahiert alle unterstützten Dokumente aus einem Verzeichnis (rekursiv).
    # Unterstützt: PDF, DOCX, XLSX.
    all_docs = []
    for file_path in folder_path.rglob("*"):
        if file_path.suffix.lower() == ".pdf":
            all_docs.extend(extract_text_from_pdf(file_path))
        elif file_path.suffix.lower() == ".docx":
            all_docs.extend(extract_text_from_docx(file_path))
        elif file_path.suffix.lower() == ".xlsx":
            all_docs.extend(extract_text_from_xlsx(file_path))
    return all_docs
